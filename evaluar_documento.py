import docx
import json

# Leer el documento
doc = docx.Document('Eva 2 Proyecto de titulo.docx')

# Extraer todo el texto
full_text = '\n'.join([p.text for p in doc.paragraphs])

# Guardar en archivo para análisis
with open('documento_completo.txt', 'w', encoding='utf-8') as f:
    f.write(full_text)

# Buscar secciones clave
sections = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and len(text) < 150:
        # Detectar títulos potenciales
        if text.isupper() or (p.style.name and 'Heading' in p.style.name):
            sections.append(f"{i}: {text}")

print(f"Total de párrafos: {len(doc.paragraphs)}")
print(f"Total de caracteres: {len(full_text)}")
print("\n=== SECCIONES DETECTADAS ===")
for section in sections[:50]:
    print(section)

# Buscar criterios de la rúbrica
criterios_buscar = [
    "análisis comparativo",
    "herramientas",
    "lenguaje",
    "componentes de hardware",
    "wireframe",
    "modelo de datos",
    "topología",
    "arquitectura",
    "SMART",
    "SLA",
    "plan de pruebas",
    "procedimientos",
    "disponibilidad",
    "casos de uso"
]

print("\n=== BÚSQUEDA DE CRITERIOS ===")
for criterio in criterios_buscar:
    count = full_text.lower().count(criterio.lower())
    if count > 0:
        print(f"✓ '{criterio}': encontrado {count} veces")
    else:
        print(f"✗ '{criterio}': NO encontrado")
